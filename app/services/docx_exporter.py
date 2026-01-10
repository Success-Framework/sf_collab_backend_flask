import os
from datetime import datetime
from docx import Document

BASE_DIR = os.path.abspath(os.path.join(os.path.dirname(__file__), '../..'))
EXPORT_DIR = os.path.join(BASE_DIR, 'uploads', 'business_plans')
os.makedirs(EXPORT_DIR, exist_ok=True)

SECTION_ORDER = [
    "executive_summary",
    "problem_solution",
    "market_overview",
    "business_model",
    "operations",
    "marketing",
    "financial_narrative"
]


def export_business_plan_docx(plan, sections, financials, projections=None):
    filename = f"business_plan_{plan.id}_{datetime.utcnow().strftime('%Y%m%d_%H%M%S')}.docx"
    file_path = os.path.join(EXPORT_DIR, filename)

    doc = Document()

    # ---------------- COVER PAGE ----------------
    doc.add_heading(plan.title, level=0)
    doc.add_paragraph(f"Industry: {plan.industry}")
    doc.add_paragraph(f"Stage: {plan.stage}")
    doc.add_paragraph(f"Generated on {datetime.utcnow().strftime('%Y-%m-%d')}")
    doc.add_page_break()

    # ---------------- TABLE OF CONTENTS (STATIC) ----------------
    doc.add_heading("Table of Contents", level=1)
    for title in [
        "Executive Summary",
        "Problem & Solution",
        "Market Overview",
        "Business Model",
        "Operations Plan",
        "Marketing Strategy",
        "Financial Overview"
    ]:
        doc.add_paragraph(title)
    doc.add_page_break()

    # ---------------- SECTIONS ----------------
    section_map = {s.section_type: s for s in sections}

    for section_type in SECTION_ORDER:
        section = section_map.get(section_type)
        if not section:
            continue

        title = section_type.replace("_", " ").title()
        doc.add_heading(title, level=1)

        for line in section.content.split("\n"):
            doc.add_paragraph(line)

        doc.add_paragraph("")

    # ---------------- FINANCIALS ----------------
    if projections:
        doc.add_page_break()
        doc.add_heading("Financial Overview", level=1)

        table = doc.add_table(rows=1, cols=4)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Year"
        hdr_cells[1].text = "Revenue"
        hdr_cells[2].text = "Expenses"
        hdr_cells[3].text = "Profit / Loss"

        for p in projections:
            row = table.add_row().cells
            row[0].text = str(p["year"])
            row[1].text = f"${p['revenue']}"
            row[2].text = f"${p['expenses']}"
            row[3].text = f"${p['profit']}"

    doc.save(file_path)
    return file_path
